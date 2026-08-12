"""Text extraction from uploaded files (PDF, DOCX, TXT)."""

import io
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument


class ParserService:
    """Extracts plain text from various document formats."""

    @staticmethod
    async def extract_text(file_path: str | None = None, file_bytes: bytes | None = None,
                           content_type: str = "text/plain") -> str:
        """Extract text from a file path or raw bytes.

        Args:
            file_path: Path to the file on disk.
            file_bytes: Raw file bytes (alternative to file_path).
            content_type: MIME type of the file.

        Returns:
            Extracted plain text string.
        """
        if file_path:
            path = Path(file_path)
            suffix = path.suffix.lower()

            if suffix == ".pdf" or content_type == "application/pdf":
                return ParserService._extract_pdf(file_path=str(path))
            elif suffix == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return ParserService._extract_docx(file_path=str(path))
            else:
                # Treat as plain text
                return path.read_text(encoding="utf-8", errors="replace")

        elif file_bytes:
            if content_type == "application/pdf":
                return ParserService._extract_pdf(file_bytes=file_bytes)
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return ParserService._extract_docx(file_bytes=file_bytes)
            else:
                return file_bytes.decode("utf-8", errors="replace")

        return ""

    @staticmethod
    def _extract_pdf(file_path: str | None = None, file_bytes: bytes | None = None) -> str:
        """Extract text from a PDF using PyMuPDF."""
        if file_path:
            doc = fitz.open(file_path)
        else:
            doc = fitz.open(stream=file_bytes, filetype="pdf")

        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()

        return "\n\n".join(text_parts).strip()

    @staticmethod
    def _extract_docx(file_path: str | None = None, file_bytes: bytes | None = None) -> str:
        """Extract text from a DOCX using python-docx."""
        if file_path:
            doc = DocxDocument(file_path)
        else:
            doc = DocxDocument(io.BytesIO(file_bytes))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs).strip()
